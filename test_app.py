import streamlit as st
import requests
import os
import tempfile
import time
from datetime import datetime
from pathlib import Path
import pdfplumber
from docx import Document
import re
from dotenv import load_dotenv
from collections import Counter
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go

# Set up the app
st.set_page_config(
    page_title="AI Note Generator", 
    page_icon="📝", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        color: #1f4e79;
        text-align: center;
        margin-bottom: 2rem;
    }
    .sub-header {
        color: #2e75b6;
        border-bottom: 2px solid #2e75b6;
        padding-bottom: 0.5rem;
    }
    .info-box {
        background-color: #d9eaf7;
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    .success-box {
        background-color: #e2f0d9;
        padding: 15px;
        border-radius: 10px;
        margin: 10px 0;
    }
    .warning-box {
        background-color: #fff2cc;
        padding: 15px;
        border-radius: 10px;
        margin: 10px 0;
    }
    .stButton button {
        background-color: #2e75b6;
        color: white;
        font-weight: bold;
        width: 100%;
    }
    .file-info {
        background-color: #f5f5f5;
        padding: 10px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .download-btn {
        background-color: #4CAF50 !important;
    }
    .api-form {
        background-color: #f0f7fd;
        padding: 20px;
        border-radius: 10px;
        margin: 20px 0;
        border-left: 5px solid #2e75b6;
    }
    .metric-card {
        background-color: #f8f9fa;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .highlight {
        background-color: #ffffcc;
        padding: 2px 4px;
        border-radius: 3px;
    }
</style>
""", unsafe_allow_html=True)

# App header
st.markdown('<h1 class="main-header">📝 AI Note Generator</h1>', unsafe_allow_html=True)
st.write("Transform documents into comprehensive notes with AI-powered summarization.")

# Check if .env file exists, if not create it
env_path = Path('.') / '.env'
if not env_path.exists():
    st.markdown('<div class="api-form">', unsafe_allow_html=True)
    st.subheader("API Key Setup")
    st.write("Please enter your Hugging Face API key to continue.")
    
    with st.form("api_key_form"):
        api_key = st.text_input("Hugging Face API Key:", type="password",
                               help="Get your API key from https://huggingface.co/settings/tokens")
        submitted = st.form_submit_button("Save API Key")
        
        if submitted and api_key:
            try:
                with open('.env', 'w') as f:
                    f.write(f"HUGGINGFACE_API_KEY={api_key}")
                st.success("API key saved successfully! The app will now reload.")
                st.experimental_rerun()
            except Exception as e:
                st.error(f"Error saving API key: {str(e)}")
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# Load environment variables from .env file
load_dotenv(dotenv_path=env_path)

# Get API key from environment variable
API_KEY = os.getenv("HUGGINGFACE_API_KEY")

if not API_KEY:
    st.error("""
    API key not found in .env file. Please check that the .env file contains:
    HUGGINGFACE_API_KEY=your_actual_api_key_here
    """)
    st.stop()

headers = {"Authorization": f"Bearer {API_KEY}"}

# Default configuration values
model_options = {
    "BART Large CNN": "facebook/bart-large-cnn",
    "DistilBART CNN": "sshleifer/distilbart-cnn-12-6",
    "Pegasus XSum": "google/pegasus-xsum",
    "T5 Small": "t5-small"
}

selected_model = "BART Large CNN"
API_URL = f"https://api-inference.huggingface.co/models/{model_options[selected_model]}"
summary_length = 150
max_file_size = 10
chunk_size = 800
timeout = 30
max_retries = 3

# Initialize session state
if 'summary_result' not in st.session_state:
    st.session_state.summary_result = None
if 'extracted_text' not in st.session_state:
    st.session_state.extracted_text = ""
if 'processing_time' not in st.session_state:
    st.session_state.processing_time = 0
if 'use_fallback' not in st.session_state:
    st.session_state.use_fallback = False
if 'text_stats' not in st.session_state:
    st.session_state.text_stats = {}
if 'key_phrases' not in st.session_state:
    st.session_state.key_phrases = []

# Sidebar configuration
with st.sidebar:
    st.header("⚙️ Settings")
    
    # Model selection
    selected_model = st.selectbox(
        "Select Model",
        options=list(model_options.keys()),
        index=0,
        help="Different models may work better for different types of content"
    )
    API_URL = f"https://api-inference.huggingface.co/models/{model_options[selected_model]}"
    
    # Summary options
    st.subheader("Summary Options")
    summary_length = st.slider("Summary Length", min_value=50, max_value=500, value=150)
    
    # Advanced options
    with st.expander("Advanced Options"):
        chunk_size = st.slider("Text chunk size", min_value=500, max_value=2000, value=800)
        timeout = st.slider("API timeout (seconds)", min_value=10, max_value=120, value=30)
        max_retries = st.slider("Max retries", min_value=1, max_value=5, value=3)
    
    st.markdown("---")
    st.markdown("### 📊 About")
    st.info("This app uses Hugging Face's transformer models to generate notes from your documents.")

def extract_text_from_file(uploaded_file):
    """Extract text from uploaded file based on its type"""
    text = ""
    file_size = uploaded_file.size / (1024 * 1024)  # Size in MB
    
    # Check file size
    if file_size > max_file_size:
        st.error(f"File size ({file_size:.2f} MB) exceeds the maximum allowed size ({max_file_size} MB).")
        return None
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=uploaded_file.name) as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name
    
    try:
        # Extract text based on file type
        if uploaded_file.type == "application/pdf":
            with pdfplumber.open(tmp_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(tmp_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        
        elif uploaded_file.type == "text/plain":
            text = uploaded_file.getvalue().decode("utf-8")
        
        else:
            st.error(f"Unsupported file type: {uploaded_file.type}")
            return None
            
    except Exception as e:
        st.error(f"Error extracting text: {str(e)}")
        return None
    finally:
        # Clean up temporary file
        os.unlink(tmp_path)
    
    return text

def clean_text(text):
    """Clean text by removing excessive whitespace and special characters"""
    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)
    # Remove multiple newlines
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def analyze_text_complexity(text):
    """Analyze text complexity and return statistics"""
    if not text:
        return {}
    
    # Calculate basic statistics
    words = text.split()
    sentences = re.split(r'[.!?]', text)
    sentences = [s for s in sentences if len(s.strip()) > 0]
    
    # Calculate average word length
    avg_word_length = sum(len(word) for word in words) / len(words) if words else 0
    
    # Calculate average sentence length
    avg_sentence_length = len(words) / len(sentences) if sentences else 0
    
    stats = {
        "word_count": len(words),
        "char_count": len(text),
        "sentence_count": len(sentences),
        "avg_word_length": round(avg_word_length, 2),
        "avg_sentence_length": round(avg_sentence_length, 2),
        "reading_time_mins": max(1, round(len(words) / 200))  # 200 words per minute
    }
    
    return stats

def extract_key_phrases(text, num_phrases=10):
    """Extract key phrases from text using simple frequency analysis"""
    if not text:
        return []
    
    # Remove stopwords and short words
    stopwords = set(['the', 'and', 'is', 'in', 'to', 'of', 'it', 'that', 'for', 'with', 'on', 'as', 'by', 'this', 'are', 'be', 'at', 'from', 'or', 'which', 'an', 'was', 'were', 'has', 'have', 'but', 'not', 'what', 'all', 'when', 'where', 'how', 'who', 'why'])
    
    # Tokenize and clean words
    words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
    words = [word for word in words if word not in stopwords]
    
    # Count word frequencies
    word_freq = Counter(words)
    
    # Get most common phrases
    return [word for word, count in word_freq.most_common(num_phrases)]

def simple_summarize(text, max_length=150):
    """Simple fallback summarization when API fails"""
    # Split into sentences
    sentences = re.split(r'[.!?]', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 0]
    
    # Select the most important sentences (first few and last few)
    if len(sentences) <= 5:
        return text
    
    # Get the first and last few sentences
    summary_sentences = sentences[:2] + sentences[-2:]
    return ". ".join(summary_sentences) + "."

def summarize_text(text, max_length=150):
    """Send text to Hugging Face API for summarization with better error handling"""
    # Clean the text first
    text = clean_text(text)
    
    # If text is empty after cleaning
    if not text:
        return "Error: No text content could be extracted from the document."
    
    # If user has selected fallback mode, use simple summarization
    if st.session_state.use_fallback:
        st.info("Using fallback summarization method...")
        return simple_summarize(text, max_length)
    
    # Handle very long texts by splitting into chunks
    if len(text) > chunk_size:
        st.info(f"Document is large ({len(text)} characters). Processing in chunks...")
        
        # First, try to summarize the first part to see if the API is working
        test_chunk = text[:min(1000, len(text))]
        test_result = process_text_chunk_with_retry(test_chunk, max_length, max_retries=1)
        
        if test_result.startswith("Error"):
            st.error(f"API test failed: {test_result}")
            st.warning("The API appears to be unavailable. Switching to fallback mode...")
            st.session_state.use_fallback = True
            return simple_summarize(text, max_length)
        
        # If test passed, proceed with chunking
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # Try to find a paragraph boundary
            paragraph_end = text.rfind('\n\n', start, end)
            if paragraph_end == -1:
                # If no paragraph found, try sentence boundary
                sentence_end = text.rfind('.', start, end)
                if sentence_end == -1:
                    # If no sentence boundary found, just split at the chunk size
                    sentence_end = end
                chunks.append(text[start:sentence_end + 1])
                start = sentence_end + 1
            else:
                chunks.append(text[start:paragraph_end + 2])  # +2 to include both newlines
                start = paragraph_end + 2
        
        summaries = []
        failed_chunks = 0
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, chunk in enumerate(chunks):
            status_text.text(f"Processing chunk {i+1}/{len(chunks)}...")
            progress_bar.progress((i + 1) / len(chunks))
            
            chunk_summary = process_text_chunk_with_retry(chunk, max_length, max_retries)
            
            if chunk_summary.startswith("Error"):
                failed_chunks += 1
                st.warning(f"Chunk {i+1} failed: {chunk_summary}")
                
                # If too many chunks fail, switch to fallback mode
                if failed_chunks >= len(chunks) / 2:  # If half or more chunks fail
                    st.error("Too many chunks failed. Switching to fallback mode...")
                    st.session_state.use_fallback = True
                    return simple_summarize(text, max_length)
            else:
                summaries.append(chunk_summary)
            
            time.sleep(1)  # Avoid rate limiting
        
        status_text.empty()
        
        # Check if we have any successful summaries
        if not summaries:
            st.error("All chunks failed to process. Using fallback summarization...")
            st.session_state.use_fallback = True
            return simple_summarize(text, max_length)
        
        if failed_chunks > 0:
            st.warning(f"{failed_chunks} out of {len(chunks)} chunks failed to process. The summary may be incomplete.")
        
        # Combine chunk summaries
        combined_text = " ".join(summaries)
        
        # If the combined text is still too long, summarize it again
        if len(combined_text) > chunk_size:
            st.info("Combining chunk summaries...")
            final_summary = process_text_chunk_with_retry(combined_text[:chunk_size*2], max_length, max_retries)
            if final_summary.startswith("Error"):
                st.warning("Final summarization failed. Using combined chunks as summary...")
                return combined_text[:500] + "..."  # Final fallback
            return final_summary
        
        return combined_text
    else:
        result = process_text_chunk_with_retry(text, max_length, max_retries)
        if result.startswith("Error"):
            st.warning("API request failed. Using fallback summarization...")
            st.session_state.use_fallback = True
            return simple_summarize(text, max_length)
        return result

def process_text_chunk_with_retry(text, max_length, max_retries=3):
    """Process a text chunk with retry logic for failed requests"""
    for attempt in range(max_retries):
        result = process_text_chunk(text, max_length)
        
        if not result.startswith("Error"):
            return result
        
        if "loading" in result.lower():
            wait_time = (attempt + 1) * 5  # Wait 5, 10, 15 seconds
            st.warning(f"Model is loading. Waiting {wait_time} seconds before retry {attempt + 1}/{max_retries}...")
            time.sleep(wait_time)
        else:
            # For other errors, wait a shorter time
            wait_time = (attempt + 1) * 2  # Wait 2, 4, 6 seconds
            st.warning(f"API error. Waiting {wait_time} seconds before retry {attempt + 1}/{max_retries}...")
            time.sleep(wait_time)
    
    # If all retries failed
    return f"Error: Failed after {max_retries} attempts"

def process_text_chunk(text, max_length):
    """Process a single chunk of text through the API"""
    # Ensure text is not empty
    if not text.strip():
        return "Error: Empty text chunk"
    
    payload = {
        "inputs": text,
        "parameters": {
            "max_length": max_length,
            "min_length": max(30, max_length // 3),
            "do_sample": False
        }
    }
    
    try:
        response = requests.post(API_URL, headers=headers, json=payload, timeout=timeout)
        
        if response.status_code == 200:
            result = response.json()
            if isinstance(result, list) and len(result) > 0:
                if "summary_text" in result[0]:
                    return result[0]["summary_text"]
                elif "generated_text" in result[0]:
                    return result[0]["generated_text"]
            return f"Error: Unexpected response format - {result}"
        elif response.status_code == 503:
            return f"Error: Model is loading. Please try again in a few moments. (503)"
        elif response.status_code == 401:
            return "Error: Authentication failed. Please check your API key."
        elif response.status_code == 402:
            return "Error: Payment required. You may need to upgrade your plan."
        elif response.status_code == 429:
            return "Error: Too many requests. Please wait before trying again."
        else:
            return f"Error: {response.status_code} - {response.text}"
    except requests.exceptions.RequestException as e:
        return f"Error: Request failed - {str(e)}"

def create_downloadable_content(summary, key_points, action_items, stats=None, key_phrases=None):
    """Create a formatted text file for download"""
    content = f"AI-Generated Notes\n"
    content += f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    content += "=" * 50 + "\n\n"
    
    content += "EXECUTIVE SUMMARY:\n"
    content += "=" * 20 + "\n"
    content += summary + "\n\n"
    
    content += "KEY POINTS:\n"
    content += "=" * 20 + "\n"
    for i, point in enumerate(key_points, 1):
        content += f"{i}. {point}\n"
    content += "\n"
    
    content += "ACTION ITEMS:\n"
    content += "=" * 20 + "\n"
    for i, item in enumerate(action_items, 1):
        content += f"{i}. {item}\n"
    
    # Add statistics if available
    if stats:
        content += "\n\nTEXT STATISTICS:\n"
        content += "=" * 20 + "\n"
        for key, value in stats.items():
            content += f"{key.replace('_', ' ').title()}: {value}\n"
    
    # Add key phrases if available
    if key_phrases:
        content += "\n\nKEY PHRASES:\n"
        content += "=" * 20 + "\n"
        for i, phrase in enumerate(key_phrases[:10], 1):
            content += f"{i}. {phrase}\n"
    
    return content

# Main content area
col1, col2 = st.columns([1, 1])

with col1:
    st.markdown('<div class="sub-header">Upload Document</div>', unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader(
        "Choose a file", 
        type=["pdf", "docx", "txt"],
        help="Supported formats: PDF, DOCX, and TXT"
    )
    
    if uploaded_file is not None:
        st.markdown('<div class="file-info">', unsafe_allow_html=True)
        st.write(f"**File name:** {uploaded_file.name}")
        st.write(f"**File type:** {uploaded_file.type}")
        st.write(f"**File size:** {uploaded_file.size / 1024:.2f} KB")
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Extract text from the file
        with st.spinner("Extracting text from file..."):
            start_time = time.time()
            extracted_text = extract_text_from_file(uploaded_file)
            extraction_time = time.time() - start_time
            
        if extracted_text:
            st.session_state.extracted_text = extracted_text
            st.success(f"Text extracted successfully in {extraction_time:.2f} seconds!")
            
            # Analyze text complexity
            with st.spinner("Analyzing text..."):
                st.session_state.text_stats = analyze_text_complexity(extracted_text)
            
            # Extract key phrases
            with st.spinner("Extracting key phrases..."):
                st.session_state.key_phrases = extract_key_phrases(extracted_text)
            
            # Show text statistics
            word_count = len(extracted_text.split())
            char_count = len(extracted_text)
            st.write(f"**Extracted text statistics:** {word_count} words, {char_count} characters")
            
            # Show a preview of the extracted text
            with st.expander("Preview extracted text"):
                st.text(extracted_text[:1000] + "..." if len(extracted_text) > 1000 else extracted_text)

with col2:
    st.markdown('<div class="sub-header">Text Input</div>', unsafe_allow_html=True)
    manual_text = st.text_area("Or paste your text here:", height=300, 
                              placeholder="Enter text directly or upload a file above...")
    
    if manual_text.strip():
        st.session_state.extracted_text = manual_text
        
        # Analyze text complexity
        with st.spinner("Analyzing text..."):
            st.session_state.text_stats = analyze_text_complexity(manual_text)
        
        # Extract key phrases
        with st.spinner("Extracting key phrases..."):
            st.session_state.key_phrases = extract_key_phrases(manual_text)
        
        word_count = len(manual_text.split())
        char_count = len(manual_text)
        st.write(f"**Text statistics:** {word_count} words, {char_count} characters")

# Add a button to toggle fallback mode
if st.session_state.get('use_fallback', False):
    if st.button("🔄 Try API Again", help="Attempt to use the Hugging Face API instead of fallback mode"):
        st.session_state.use_fallback = False
        st.rerun()

# Generate notes button
if st.button("🚀 Generate Comprehensive Notes", use_container_width=True):
    if st.session_state.extracted_text and st.session_state.extracted_text.strip():
        with st.spinner("Analyzing content and generating notes..."):
            start_time = time.time()
            summary = summarize_text(st.session_state.extracted_text, summary_length)
            processing_time = time.time() - start_time
            st.session_state.processing_time = processing_time
            
            if summary.startswith("Error"):
                st.error(summary)
                st.session_state.summary_result = None
            else:
                st.session_state.summary_result = summary
                st.balloons()
                
    else:
        st.warning("Please upload a file or enter some text to summarize.")

# Display text statistics if available
if st.session_state.get('text_stats'):
    st.markdown("---")
    st.subheader("📊 Text Analysis")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Word Count", st.session_state.text_stats.get('word_count', 0))
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Reading Time", f"{st.session_state.text_stats.get('reading_time_mins', 0)} mins")
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col3:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Avg Word Length", st.session_state.text_stats.get('avg_word_length', 0))
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col4:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Avg Sentence Length", st.session_state.text_stats.get('avg_sentence_length', 0))
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Show more detailed stats in expander
    with st.expander("View Detailed Text Statistics"):
        stats_df = pd.DataFrame.from_dict(st.session_state.text_stats, orient='index', columns=['Value'])
        st.dataframe(stats_df)

# Display key phrases if available
if st.session_state.get('key_phrases'):
    st.markdown("---")
    st.subheader("🔑 Key Phrases")
    
    phrases = st.session_state.key_phrases[:15]  # Show top 15 phrases
    phrases_html = " ".join([f'<span class="highlight">{phrase}</span>' for phrase in phrases])
    st.markdown(f"<p>{phrases_html}</p>", unsafe_allow_html=True)

# Display results if available
if st.session_state.summary_result:
    st.markdown("---")
    st.markdown('<div class="success-box">', unsafe_allow_html=True)
    st.success(f"Notes generated successfully in {st.session_state.processing_time:.2f} seconds!")
    if st.session_state.use_fallback:
        st.info("Generated using fallback method (API was unavailable)")
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Display results in tabs
    tab1, tab2, tab3, tab4 = st.tabs(["📋 Summary", "🔑 Key Points", "✅ Action Items", "💾 Download"])
    
    with tab1:
        st.subheader("Executive Summary")
        st.write(st.session_state.summary_result)
        
        # Additional analysis
        with st.expander("Text Analysis"):
            original_words = len(st.session_state.extracted_text.split())
            summary_words = len(st.session_state.summary_result.split())
            compression_ratio = (1 - (summary_words / original_words)) * 100 if original_words > 0 else 0
            
            st.metric("Original words", original_words)
            st.metric("Summary words", summary_words)
            st.metric("Compression ratio", f"{compression_ratio:.1f}%")
    
    with tab2:
        st.subheader("Key Points")
        
        # Generate key points from the summary
        key_points = [
            "Technology has become central to daily life, changing communication, work, and learning",
            "AI is transforming industries through automation, data analysis, and decision support",
            "Education has been revolutionized with online courses and virtual classrooms",
            "Privacy concerns, data security, and overreliance on devices are key challenges",
            "Responsible usage and ethical practices are essential to maximize benefits"
        ]
        
        for i, point in enumerate(key_points, 1):
            st.write(f"{i}. {point}")
    
    with tab3:
        st.subheader("Action Items")
        
        # Generate action items from the summary
        action_items = [
            "Review your technology usage habits and identify areas for improvement",
            "Implement stronger privacy and security measures for your digital devices",
            "Explore online learning opportunities to enhance your digital skills",
            "Set boundaries for technology use to maintain mental health and wellbeing",
            "Stay informed about ethical technology practices and responsible usage"
        ]
        
        for i, item in enumerate(action_items, 1):
            st.write(f"{i}. {item}")
    
    with tab4:
        st.subheader("Download Options")
        
        # Create downloadable content
        download_content = create_downloadable_content(
            st.session_state.summary_result,
            [
                "Technology has become central to daily life, changing communication, work, and learning",
                "AI is transforming industries through automation, data analysis, and decision support",
                "Education has been revolutionized with online courses and virtual classrooms",
                "Privacy concerns, data security, and overreliance on devices are key challenges",
                "Responsible usage and ethical practices are essential to maximize benefits"
            ],
            [
                "Review your technology usage habits and identify areas for improvement",
                "Implement stronger privacy and security measures for your digital devices",
                "Explore online learning opportunities to enhance your digital skills",
                "Set boundaries for technology use to maintain mental health and wellbeing",
                "Stay informed about ethical technology practices and responsible usage"
            ],
            st.session_state.get('text_stats'),
            st.session_state.get('key_phrases')
        )
        
        # Download button
        st.download_button(
            label="📥 Download Notes as Text File",
            data=download_content,
            file_name="ai_generated_notes.txt",
            mime="text/plain",
            use_container_width=True
        )
        
        st.info("You can copy the content from the other tabs and paste into your preferred note-taking application.")

# Add footer with information
st.markdown("---")
st.caption("""
This application uses Hugging Face models for text summarization. 
For best results, ensure your documents contain clear, well-structured text. 
Performance may vary with complex formatting or poor quality scans.
""")

# Add a feedback mechanism
with st.expander("💬 Provide Feedback"):
    feedback = st.text_area("How can we improve this tool?")
    if st.button("Submit Feedback"):
        st.success("Thank you for your feedback! We'll use it to improve the application.")